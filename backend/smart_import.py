"""
智能材料导入系统 - 核心流水线
自动识别、分类、归档各类材料文件
"""

import os
import json
import uuid
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime, date
from difflib import SequenceMatcher

from fastapi import UploadFile
from PIL import Image

from database import (
    get_session, Company, Person, Material, Document, PendingReview, MaterialVersion
)
from ocr_client import ocr_image, check_ocr_service
from llm_provider import get_llm_provider
from ocr_cache import get_cached_ocr, save_ocr_to_cache, get_file_hash
from page_extraction_strategy import get_pages_to_extract
from contract_page_analyzer import analyze_contract_key_pages
from certificate_matcher import find_existing_certificate, is_newer_version

logger = logging.getLogger("materialhub.smart_import")

# 临时文件目录
TEMP_DIR = Path(os.getenv("DATA_DIR", "data")) / "temp"
TEMP_DIR.mkdir(parents=True, exist_ok=True)


def extract_pdf_page_to_png(pdf_path: str, page_num: int, output_dir: Path, dpi: int = 300) -> Optional[str]:
    """
    从PDF提取指定页面并转换为PNG图片

    Args:
        pdf_path: PDF文件路径
        page_num: 页码（从0开始，-1表示最后一页）
        output_dir: 输出目录
        dpi: 分辨率（默认300）

    Returns:
        PNG文件路径，失败返回None
    """
    import fitz  # PyMuPDF

    try:
        doc = fitz.open(pdf_path)

        # 处理负数索引（-1表示最后一页）
        if page_num < 0:
            page_num = len(doc) + page_num

        # 检查页码是否有效
        if page_num < 0 or page_num >= len(doc):
            logger.error(f"页码无效: {page_num}, 文档总页数: {len(doc)}")
            return None

        page = doc[page_num]

        # 渲染为高分辨率图片
        mat = fitz.Matrix(dpi/72, dpi/72)
        pix = page.get_pixmap(matrix=mat)

        # 保存为PNG
        output_path = output_dir / f"page_{page_num + 1}_{uuid.uuid4().hex[:8]}.png"
        pix.save(str(output_path))

        doc.close()

        logger.info(f"  ✓ 提取页面{page_num + 1}为PNG: {output_path.name}")
        return str(output_path)

    except Exception as e:
        logger.error(f"提取PDF页面失败: {e}")
        return None


class FileProcessor:
    """多格式文件处理器"""

    SUPPORTED_FORMATS = {
        'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'],
        'document': ['.docx', '.doc', '.pdf'],
        'other': ['.txt', '.xlsx']
    }

    def analyze_file(self, file: UploadFile) -> Dict:
        """分析文件类型和特征"""
        filename = file.filename.lower()
        ext = Path(filename).suffix

        # 检测文件类型
        if ext in self.SUPPORTED_FORMATS['image']:
            file_type = 'image'
        elif ext in self.SUPPORTED_FORMATS['document']:
            file_type = 'document'
        else:
            file_type = 'other'

        # 从文件名猜测材料类型
        filename_hints = self.extract_hints_from_filename(filename)

        return {
            "type": file_type,
            "extension": ext,
            "filename": file.filename,
            "hints": filename_hints,
        }

    def extract_hints_from_filename(self, filename: str) -> Dict:
        """从文件名提取提示信息"""
        import re

        hints = {
            "possible_company": None,
            "possible_type": None,
            "possible_date": None
        }

        # 识别公司名（常见模式）
        company_patterns = [
            r'([\u4e00-\u9fa5]+(?:科技|网络|信息|技术|系统|软件|数据|智能)[\u4e00-\u9fa5]*(?:有限公司|公司|集团))',
            r'([\u4e00-\u9fa5]{2,20}有限公司)',
        ]

        for pattern in company_patterns:
            match = re.search(pattern, filename)
            if match:
                hints["possible_company"] = match.group(1)
                break

        # 识别材料类型关键词
        type_keywords = {
            "营业执照": "license",
            "资质": "qualification",
            "ISO": "iso_cert",
            "身份证": "id_card",
            "法人": "legal_person_cert",
            "毕业证": "education_cert",
            "学历": "education_cert",
            "完税证明": "tax_payment_cert",
            "纳税证明": "tax_payment_cert",
            "完税": "tax_payment_cert",
            "交税凭证": "tax_payment_voucher",
            "缴税凭证": "tax_payment_voucher",
            "缴税": "tax_payment_voucher",
            "审计报告": "audit_report",
            "审计": "audit_report",
            "财务报告": "audit_report",
        }

        for keyword, mat_type in type_keywords.items():
            if keyword in filename:
                hints["possible_type"] = mat_type
                break

        # 提取日期
        date_match = re.search(r'(20\d{2})[年\-_]?(\d{1,2})?[月\-_]?(\d{1,2})?', filename)
        if date_match:
            hints["possible_date"] = date_match.group(0)

        return hints


class ContentExtractor:
    """统一的内容提取器"""

    async def extract(self, file_path: str, file_info: Dict, progress_callback=None, start_page: int = 0, page_limit: int = None, page_numbers: List[int] = None) -> Dict:
        """根据文件类型提取内容

        参数:
        - page_numbers: 指定要扫描的页码列表（从1开始），例如：[1, 3, 5]
                       如果提供此参数，start_page 和 page_limit 将被忽略
        """

        if file_info["type"] == "image":
            return await self.extract_from_image(file_path)

        elif file_info["type"] == "document":
            if file_info["extension"] == ".pdf":
                return await self.extract_from_pdf(file_path, progress_callback=progress_callback, start_page=start_page, page_limit=page_limit, page_numbers=page_numbers)
            elif file_info["extension"] in [".docx", ".doc"]:
                return await self.extract_from_docx(file_path)

        return {"text": "", "images": [], "metadata": {}}

    async def extract_from_image(self, file_path: str) -> Dict:
        """从图片提取文本"""
        if not check_ocr_service():
            logger.warning("OCR服务不可用")
            return {"text": "", "images": [file_path], "metadata": {"source": "no_ocr"}}

        ocr_text = ocr_image(file_path)

        return {
            "text": ocr_text or "",
            "images": [file_path],
            "metadata": {
                "image_count": 1,
                "source": "ocr"
            }
        }

    async def extract_from_pdf(self, file_path: str, progress_callback=None, start_page: int = 0, page_limit: int = None, page_numbers: List[int] = None) -> Dict:
        """从PDF提取文本和图片"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF未安装，无法处理PDF")
            return {"text": "", "images": [], "metadata": {"error": "PyMuPDF not installed"}}

        doc = fitz.open(file_path)

        # 提取文本
        text_content = []
        images = []

        for page_num in range(len(doc)):
            page = doc[page_num]

            # 提取文本
            text = page.get_text()
            text_content.append(text)

            # 提取嵌入的图片
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)

                # 保存图片
                img_path = TEMP_DIR / f"pdf_img_{uuid.uuid4()}.{base_image['ext']}"
                with open(img_path, "wb") as img_file:
                    img_file.write(base_image["image"])

                images.append(str(img_path))

        full_text = "\n".join(text_content)

        # 如果PDF是扫描件（文本很少），将PDF页面渲染为图片并OCR
        if len(full_text.strip()) < 100 and check_ocr_service():
            total_pages = len(doc)

            # 确定扫描范围
            if page_numbers:
                # 用户指定页码 - 只扫描指定页面
                # page_numbers 是从1开始的，需要转换为从0开始的索引
                pages_to_scan_indices = [p - 1 for p in page_numbers if 1 <= p <= total_pages]
                logger.info(f"检测到扫描PDF，共{total_pages}页，用户指定扫描页码: {page_numbers}（共{len(pages_to_scan_indices)}页）")
            else:
                # 默认范围扫描模式
                # start_page: 开始页码（从0开始）
                # page_limit: 最多扫描多少页（None表示不限制）
                actual_start = start_page
                if page_limit:
                    actual_end = min(start_page + page_limit, total_pages, 20)  # 最多20页
                else:
                    actual_end = min(20, total_pages)

                pages_to_scan_indices = list(range(actual_start, actual_end))

                if start_page == 0:
                    logger.info(f"检测到扫描PDF，共{total_pages}页，先扫描前{len(pages_to_scan_indices)}页...")
                else:
                    logger.info(f"继续扫描PDF第{actual_start+1}-{actual_end}页，共{len(pages_to_scan_indices)}页...")

            ocr_texts = []
            ocr_results = []  # 存储每页的OCR结果预览

            if progress_callback:
                if page_numbers:
                    msg = f"开始OCR识别指定页码: {page_numbers}..."
                else:
                    msg = f"开始OCR识别..."
                progress_callback({
                    "stage": "ocr",
                    "message": msg,
                    "current_page": 0,
                    "total_pages": total_pages,
                    "ocr_results": []
                })

            for page_num in pages_to_scan_indices:
                page = doc[page_num]

                # 将页面渲染为图片 (300 DPI)
                mat = fitz.Matrix(300/72, 300/72)  # 提高分辨率以改善OCR效果
                pix = page.get_pixmap(matrix=mat)

                # 保存为临时图片
                page_img_path = TEMP_DIR / f"pdf_page_{uuid.uuid4()}.png"
                pix.save(str(page_img_path))
                images.append(str(page_img_path))

                # OCR识别（先检查缓存）
                logger.info(f"📄 OCR处理第 {page_num + 1}/{total_pages} 页...")

                if progress_callback:
                    progress_callback({
                        "stage": "ocr",
                        "message": f"正在识别第 {page_num + 1}/{total_pages} 页...",
                        "current_page": page_num + 1,
                        "total_pages": total_pages,
                        "ocr_results": ocr_results
                    })

                # 🔥 先尝试从缓存获取OCR结果
                cached_ocr = get_cached_ocr(file_path, page_num)
                if cached_ocr:
                    ocr_text = cached_ocr["text"]
                    logger.info(f"   ✨ 第{page_num + 1}页使用缓存（{len(ocr_text)}字符）")
                else:
                    # 缓存未命中，执行OCR
                    ocr_text = ocr_image(str(page_img_path))
                    if ocr_text:
                        # 保存到缓存
                        save_ocr_to_cache(file_path, page_num, ocr_text)

                if ocr_text:
                    ocr_texts.append(f"=== 第{page_num + 1}页 ===\n{ocr_text}")
                    logger.info(f"   ✅ 第{page_num + 1}页识别完成，提取{len(ocr_text)}字符")

                    # 保存OCR结果预览（前200字符）
                    preview = ocr_text[:200] + ("..." if len(ocr_text) > 200 else "")
                    ocr_results.append({
                        "page": page_num + 1,
                        "chars": len(ocr_text),
                        "preview": preview,
                        "status": "success"
                    })

                    if progress_callback:
                        progress_callback({
                            "stage": "ocr",
                            "message": f"第 {page_num + 1}/{total_pages} 页识别完成（{len(ocr_text)}字符）",
                            "current_page": page_num + 1,
                            "total_pages": total_pages,
                            "ocr_results": ocr_results
                        })
                else:
                    logger.warning(f"   ⚠️ 第{page_num + 1}页识别失败")
                    ocr_results.append({
                        "page": page_num + 1,
                        "chars": 0,
                        "preview": "",
                        "status": "failed"
                    })

            if ocr_texts:
                full_text = "\n\n".join(ocr_texts)
                if page_numbers:
                    scanned_desc = f"指定页码{page_numbers}"
                else:
                    scanned_desc = f"第{actual_start+1}-{actual_end}页"
                logger.info(f"🎉 OCR阶段完成！共处理{len(ocr_texts)}页（{scanned_desc}），提取文本 {len(full_text)} 字符")

                if progress_callback:
                    progress_callback({
                        "stage": "ocr_phase_complete",
                        "message": f"OCR阶段完成！已识别{scanned_desc}，共{len(ocr_texts)}页，提取{len(full_text)}字符",
                        "current_page": pages_to_scan_indices[-1] + 1 if pages_to_scan_indices else 0,
                        "total_pages": total_pages,
                        "ocr_results": ocr_results,
                        "scanned_pages": len(ocr_texts),
                        "total_document_pages": total_pages
                    })
            else:
                logger.warning("❌ OCR未能提取到任何文本")

        # 如果有嵌入图片但文本少，也尝试OCR嵌入图片
        elif len(full_text.strip()) < 100 and images and check_ocr_service():
            logger.info("尝试OCR嵌入图片...")
            ocr_texts = []
            for img_path in images[:10]:  # 最多10张图片
                ocr_text = ocr_image(img_path)
                if ocr_text:
                    ocr_texts.append(ocr_text)
            if ocr_texts:
                full_text = "\n".join(ocr_texts)

        # 准备返回数据的metadata
        metadata = {
            "page_count": len(doc),
            "image_count": len(images),
            "source": "pdf_extract" if len(full_text) > 100 else "ocr"
        }

        # 收集OCR结果（如果有的话）
        result = {
            "text": full_text,
            "images": images,
            "metadata": metadata
        }

        # 如果进行了OCR，添加已扫描页数信息和OCR结果
        try:
            if 'ocr_texts' in locals() and ocr_texts:
                metadata['scanned_pages'] = len(ocr_texts)
            if 'ocr_results' in locals() and ocr_results:
                result['ocr_results'] = ocr_results  # 包含每页的预览信息
        except:
            pass

        return result

    async def extract_from_docx(self, file_path: str) -> Dict:
        """从Word文档提取文本和图片"""
        from docx import Document

        doc = Document(file_path)

        # 提取文本
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # 提取图片
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                img_path = TEMP_DIR / f"docx_img_{uuid.uuid4()}.png"
                with open(img_path, "wb") as img_file:
                    img_file.write(image_data)
                images.append(str(img_path))

        return {
            "text": "\n".join(text_content),
            "images": images,
            "metadata": {
                "paragraph_count": len(doc.paragraphs),
                "image_count": len(images),
                "source": "docx_extract"
            }
        }


class IntelligentAnalyzer:
    """LLM智能分析器 - 核心大脑（两阶段智能识别）"""

    def __init__(self):
        self.llm = get_llm_provider()

    async def analyze(self, extracted_content: Dict, filename: str) -> Dict:
        """两阶段智能分析

        阶段1: 快速分类 - 识别材料类型
        阶段2: 深度提取 - 根据类型提取关键信息
        """
        text = extracted_content["text"]

        if not text or len(text.strip()) < 20:
            logger.warning(f"文本内容过少，无法进行智能分析: {filename}")
            return {
                "material_type": "unknown",
                "confidence": 0.0,
                "error": "文本内容不足"
            }

        # 🎯 阶段1: 快速分类
        material_type = await self._classify_material_type(text, filename)
        logger.info(f"📌 材料类型识别: {material_type}")

        # 🎯 阶段2: 深度提取（根据类型调用专门的提取器）
        if material_type == "contract":
            analysis = await self._extract_contract(text, filename)
        elif material_type == "company_qualification":
            analysis = await self._extract_company_qualification(text, filename)
        elif material_type == "company_business":
            analysis = await self._extract_company_business(text, filename)
        elif material_type == "iso_certificate":
            analysis = await self._extract_iso_certificate(text, filename)
        elif material_type == "employee_document":
            analysis = await self._extract_employee_document(text, filename)
        elif material_type == "project_performance":
            analysis = await self._extract_project_performance(text, filename)
        elif material_type == "financial_document":
            analysis = await self._extract_financial_document(text, filename)
        elif material_type == "tax_payment_voucher":
            analysis = await self._extract_tax_payment_voucher(text, filename)
        elif material_type == "audit_report":
            analysis = await self._extract_audit_report(text, filename)
        else:
            # 未知类型 - 使用通用提取
            analysis = await self._extract_generic(text, filename)

        # 添加原始文本
        analysis["_extracted_text"] = text
        analysis["material_type"] = material_type

        # 后处理
        analysis = self.post_process_analysis(analysis)

        return analysis

    async def _classify_material_type(self, text: str, filename: str) -> str:
        """阶段1: 快速分类材料类型（轻量级prompt）"""

        prompt = f"""你是材料分类专家。请快速判断这是什么类型的材料。

## 文件信息
文件名: {filename}

## 内容预览（前800字符）
{text[:800]}

## 材料类型分类
请从以下类型中选择最匹配的一个：

1. **contract** - 合同类（采购合同、服务合同、劳动合同、框架协议等）
   识别特征：有甲方乙方、合同编号、签订日期

2. **company_qualification** - 公司资质类（建筑资质、安全生产许可证、经营许可证等）
   识别特征：资质证书、许可证、准入证

3. **company_business** - 公司工商信息类（营业执照、法定代表人证明、公司章程等）
   识别特征：统一社会信用代码、工商注册信息

4. **iso_certificate** - ISO认证类（ISO9001、ISO14001、ISO27001等）
   识别特征：ISO标准编号、认证机构、管理体系

5. **employee_document** - 员工证件类（身份证、学历证书、职称证书、职业资格证等）
   识别特征：个人姓名、身份证号、学历、职称

6. **project_performance** - 项目业绩类（项目合同、验收报告、中标通知书等）
   识别特征：项目名称、建设单位、承建单位、项目金额

7. **financial_document** - 财务票据类（发票、收据等）
   识别特征：发票号、税号、金额、税率

8. **tax_payment_voucher** - 交税凭证类（缴税凭证、税款缴纳凭证等）
   识别特征：纳税人识别号、税款金额、缴款日期、税款所属期

9. **audit_report** - 审计报告类（财务审计报告、专项审计报告等）
   识别特征：审计报告、审计意见、会计师事务所、审计机构、财务报表

9. **other** - 其他类型

## 输出要求
只输出一个类型代码，不要有其他文字。例如：contract
"""

        try:
            response = self.llm.chat([{"role": "user", "content": prompt}])
            material_type = response.strip().lower()

            # 验证返回的类型是否合法
            valid_types = [
                "contract", "company_qualification", "company_business",
                "iso_certificate", "employee_document", "project_performance",
                "financial_document", "tax_payment_voucher", "audit_report", "other"
            ]

            if material_type not in valid_types:
                logger.warning(f"LLM返回了未知类型: {material_type}，使用other")
                material_type = "other"

            return material_type

        except Exception as e:
            logger.error(f"材料类型分类失败: {e}")
            return "other"

    async def _extract_contract(self, text: str, filename: str) -> Dict:
        """提取合同类材料信息（专门优化的prompt）"""

        # 使用字符串拼接避免f-string嵌套过深的问题
        prompt_parts = [
            "你是合同信息提取专家。请从合同中提取关键信息。\n",
            "\n## 文件名\n",
            filename,
            "\n\n## 合同内容\n",
            text[:3000],
            "\n\n## 提取任务\n",
            "请严格按以下JSON格式返回（只输出JSON，不要markdown代码块）：\n\n",
            '{\n',
            '    "material_name": "合同标准名称",\n',
            '    "contract_info": {\n',
            '        "contract_number": "合同编号（必填）",\n',
            '        "party_a": {"name": "甲方公司全称"},\n',
            '        "party_b": {"name": "乙方公司全称"},\n',
            '        "amount": "合同金额",\n',
            '        "signing_date": "签订日期（YYYY-MM-DD）"\n',
            '    },\n',
            '    "confidence": 0.90,\n',
            '    "needs_more_pages": false\n',
            '}\n\n',
            "## 核心提取规则\n\n",
            "### 1. 甲方乙方识别（最重要！）\n",
            "- 甲方通常标注：\"甲方：\"、\"需方：\"、\"采购方：\"\n",
            "- 乙方通常标注：\"乙方：\"、\"供方：\"、\"承包方：\"\n",
            "- 必须提取完整公司名称（不要简称）\n\n",
            "### 2. 合同编号识别\n",
            "- 通常在标题下方\n",
            "- 标注为：\"合同编号\"、\"协议编号\"、\"编号\"\n\n",
            "### 3. 🚨 needs_more_pages判断规则\n\n",
            "**设为 false 的条件（同时满足3项）：**\n",
            "1. 甲方公司名称不为null\n",
            "2. 乙方公司名称不为null\n",
            "3. 合同编号不为null\n\n",
            "**设为 true 的条件：**\n",
            "- 以上任一项缺失\n\n",
            "**重要提示：**\n",
            "- 金额、日期、法人等字段不影响needs_more_pages判断\n",
            "- 只要甲方+乙方+编号齐全，就设为false\n",
            "- 第一页通常就能找到这三项核心信息\n\n",
            "## 归档规则说明\n",
            "**本系统用于投标业务，乙方（party_b）是\"我们公司\"，合同将归档到乙方公司名下**\n",
            "- 甲方：客户/采购方/需方\n",
            "- 乙方：我们公司/供应商/承包方 ← 归档目标\n\n",
            "## 注意事项\n",
            "1. 未识别的字段填null\n",
            "2. 公司名称必须完整\n",
            "3. 只输出JSON，不要markdown代码块\n"
        ]

        prompt = "".join(prompt_parts)
        return await self._parse_llm_json(prompt)

    async def _extract_company_qualification(self, text: str, filename: str) -> Dict:
        """提取公司资质证书信息"""

        prompt = f"""你是资质证书信息提取专家。请从证书中提取关键信息。

## 文件名
{filename}

## 证书内容
{text[:3000]}

## 提取任务
```json
{{
    "material_name": "资质证书标准名称",
    "company_info": {{
        "name": "持证企业全称（完整）",
        "legal_person": "法定代表人",
        "credit_code": "统一社会信用代码（18位）",
        "address": "企业地址"
    }},
    "certificate_info": {{
        "certificate_number": "证书编号（必填）",
        "certificate_name": "资质名称",
        "certificate_grade": "资质等级（如：甲级、一级、二级）",
        "issuing_authority": "发证机关（必填）",
        "business_scope": "业务范围/许可范围"
    }},
    "key_dates": {{
        "issue_date": "发证日期（YYYY-MM-DD）",
        "expiry_date": "有效期至（YYYY-MM-DD）",
        "validity_period": "有效期（如：5年）"
    }},
    "confidence": 0.90,
    "needs_more_pages": false
}}
```

## 核心识别规则
1. **证书编号** - 通常在证书显著位置，格式多样
2. **发证机关** - 通常在证书底部，如"XX省住房和城乡建设厅"
3. **有效期** - 寻找"有效期至"、"有效期限"
4. **企业全称** - 必须完整，不要简称

### needs_more_pages判断
如果已识别到：证书编号 + 发证机关 + 企业名称 → 设为false
如果三要素缺失 → 设为true
"""

        return await self._parse_llm_json(prompt)

    async def _extract_company_business(self, text: str, filename: str) -> Dict:
        """提取公司工商信息（营业执照等）"""

        prompt = f"""你是工商信息提取专家。请从营业执照中提取关键信息。

## 文件名
{filename}

## 内容
{text[:3000]}

## 提取任务
```json
{{
    "material_name": "材料名称（如：营业执照）",
    "company_info": {{
        "name": "公司名称（完整全称）",
        "legal_person": "法定代表人（必填）",
        "credit_code": "统一社会信用代码（18位，必填）",
        "registered_capital": "注册资本",
        "establishment_date": "成立日期（YYYY-MM-DD）",
        "address": "注册地址",
        "business_scope": "经营范围（简要概括）",
        "company_type": "公司类型（如：有限责任公司）"
    }},
    "key_dates": {{
        "establishment_date": "成立日期（YYYY-MM-DD）",
        "issue_date": "核准日期（YYYY-MM-DD）",
        "expiry_date": "营业期限（YYYY-MM-DD，长期经营填null）"
    }},
    "confidence": 0.95,
    "needs_more_pages": false
}}
```

## 核心识别规则
1. **统一社会信用代码** - 18位，营业执照核心标识（必填）
2. **法定代表人** - 通常在显著位置（必填）
3. **公司名称** - 必须完整，包括"有限公司"等后缀

### needs_more_pages判断
如果已识别到：公司名称 + 统一社会信用代码 + 法定代表人 → 设为false
否则 → 设为true
"""

        return await self._parse_llm_json(prompt)

    async def _extract_iso_certificate(self, text: str, filename: str) -> Dict:
        """提取ISO认证证书信息"""

        prompt = f"""你是ISO认证证书信息提取专家。请从证书中提取关键信息。

## 文件名
{filename}

## 证书内容
{text[:3000]}

## 提取任务
```json
{{
    "material_name": "认证证书名称（如：ISO9001质量管理体系认证证书）",
    "company_info": {{
        "name": "获证组织全称（完整）",
        "address": "组织地址"
    }},
    "certificate_info": {{
        "certificate_number": "证书编号（必填）",
        "iso_standard": "ISO标准（如：ISO 9001:2015）",
        "issuing_authority": "认证机构（必填）",
        "certification_scope": "认证范围"
    }},
    "key_dates": {{
        "issue_date": "颁发日期（YYYY-MM-DD，必填）",
        "expiry_date": "有效期至（YYYY-MM-DD，必填）",
        "initial_certification_date": "初次认证日期（YYYY-MM-DD）"
    }},
    "confidence": 0.90,
    "needs_more_pages": false
}}
```

## 核心识别规则
1. **ISO标准编号** - 如ISO 9001:2015、ISO 14001:2015
2. **认证机构** - 通常有"认证有限公司"字样
3. **证书编号** - 认证机构颁发的唯一编号
4. **有效期** - ISO证书都有明确的有效期（通常3年）

### needs_more_pages判断
如果已识别到：证书编号 + 认证机构 + 有效期至 → 设为false
否则 → 设为true
"""

        return await self._parse_llm_json(prompt)

    async def _extract_employee_document(self, text: str, filename: str) -> Dict:
        """提取员工证件信息"""

        prompt = f"""你是员工证件信息提取专家。请提取关键信息。

## 文件名
{filename}

## 内容
{text[:3000]}

## 提取任务
```json
{{
    "material_name": "证件名称（如：身份证、本科毕业证书、高级工程师证书）",
    "person_info": {{
        "name": "姓名（必填）",
        "id_number": "身份证号（18位）",
        "gender": "性别",
        "birth_date": "出生日期（YYYY-MM-DD）",
        "education": "学历（如：本科、硕士）",
        "degree": "学位（如：学士、硕士）",
        "major": "专业",
        "university": "毕业院校",
        "position": "职位/职称",
        "certificate_type": "证书类型（学历/学位/职称/资格）"
    }},
    "certificate_info": {{
        "certificate_number": "证书编号",
        "issuing_authority": "发证机关/颁发单位"
    }},
    "key_dates": {{
        "issue_date": "颁发日期（YYYY-MM-DD）",
        "graduation_date": "毕业日期（YYYY-MM-DD）"
    }},
    "confidence": 0.90,
    "needs_more_pages": false
}}
```

## 核心识别规则
1. **身份证** - 姓名+身份证号
2. **学历证书** - 姓名+学历+毕业院校+专业
3. **职称证书** - 姓名+职称名称+发证机关

### needs_more_pages判断
如果已识别到核心字段（姓名+证件特征信息）→ 设为false
否则 → 设为true
"""

        return await self._parse_llm_json(prompt)

    async def _extract_project_performance(self, text: str, filename: str) -> Dict:
        """提取项目业绩信息"""

        prompt = f"""你是项目业绩信息提取专家。请提取关键信息。

## 文件名
{filename}

## 内容
{text[:3000]}

## 提取任务
```json
{{
    "material_name": "项目名称",
    "project_info": {{
        "project_name": "项目全称",
        "project_number": "项目编号",
        "builder": "建设单位（甲方）",
        "contractor": "承建单位（乙方）",
        "project_amount": "项目金额",
        "project_location": "项目地点",
        "project_type": "项目类型（如：建筑工程、系统集成）"
    }},
    "company_info": {{
        "name": "承建单位全称（归档目标）"
    }},
    "key_dates": {{
        "start_date": "开工日期（YYYY-MM-DD）",
        "completion_date": "竣工日期（YYYY-MM-DD）",
        "acceptance_date": "验收日期（YYYY-MM-DD）"
    }},
    "confidence": 0.85,
    "needs_more_pages": false
}}
```

## 归档说明
**承建单位（乙方）作为归档目标**
"""

        return await self._parse_llm_json(prompt)

    async def _extract_financial_document(self, text: str, filename: str) -> Dict:
        """提取财务票据信息（发票、完税证明等）"""

        prompt = f"""你是财务票据信息提取专家。请提取关键信息。

## 文件名
{filename}

## 内容
{text[:3000]}

## 提取任务
根据内容判断是**发票**还是**完税证明**，然后提取对应字段：

### 如果是发票（增值税专用发票、普通发票等）
```json
{{
    "material_name": "增值税专用发票",
    "invoice_info": {{
        "invoice_number": "发票号码",
        "invoice_code": "发票代码",
        "buyer": "购买方名称",
        "buyer_tax_number": "购买方纳税人识别号",
        "seller": "销售方名称",
        "seller_tax_number": "销售方纳税人识别号",
        "total_amount": "价税合计",
        "amount": "金额",
        "tax_amount": "税额",
        "tax_rate": "税率"
    }},
    "company_info": {{
        "name": "购买方名称（归档目标）"
    }},
    "key_dates": {{
        "issue_date": "开票日期（YYYY-MM-DD）"
    }},
    "confidence": 0.90,
    "needs_more_pages": false
}}
```

### 如果是完税证明（税收完税证明、纳税证明等）
```json
{{
    "material_name": "税收完税证明",
    "tax_payment_info": {{
        "certificate_number": "完税证明号",
        "taxpayer_name": "纳税人名称",
        "taxpayer_id": "纳税人识别号（统一社会信用代码）",
        "tax_type": "税种",
        "tax_amount": "完税金额",
        "tax_authority": "税务机关"
    }},
    "company_info": {{
        "name": "纳税人名称（归档目标）"
    }},
    "key_dates": {{
        "issue_date": "填发日期（YYYY-MM-DD）"
    }},
    "confidence": 0.95,
    "needs_more_pages": false
}}
```

## 识别规则
1. **完税证明**特征：文档标题包含"完税证明"、"纳税证明"，包含纳税人识别号、税务机关
2. **发票**特征：标题包含"发票"、"Invoice"，包含发票代码、发票号码、税率

## 归档说明
- **发票**：购买方作为归档目标
- **完税证明**：纳税人作为归档目标

## 重要提示
**只输出JSON对象，不要输出任何其他文字、说明或markdown标记。直接输出纯JSON。**
"""

        return await self._parse_llm_json(prompt)

    async def _extract_tax_payment_voucher(self, text: str, filename: str) -> Dict:
        """提取交税凭证信息"""

        prompt = f"""你是交税凭证信息提取专家。请提取关键信息。

## 文件名
{filename}

## 内容
{text[:3000]}

## 提取任务
提取交税凭证（缴税凭证、税款缴纳凭证）的关键信息：

```json
{{
    "material_name": "税款缴纳凭证/交税凭证",
    "tax_voucher_info": {{
        "voucher_number": "凭证号码/缴款书号码",
        "taxpayer_name": "纳税人名称",
        "taxpayer_id": "纳税人识别号（统一社会信用代码）",
        "tax_type": "税种（如：增值税、企业所得税等）",
        "tax_period": "税款所属期（YYYY-MM格式）",
        "payment_amount": "缴款金额",
        "payment_date": "缴款日期（YYYY-MM-DD）",
        "tax_authority": "征收机关/税务机关",
        "bank": "缴款银行"
    }},
    "company_info": {{
        "name": "纳税人名称（归档目标）"
    }},
    "key_dates": {{
        "payment_date": "缴款日期（YYYY-MM-DD）",
        "tax_period": "税款所属期（YYYY-MM）"
    }},
    "confidence": 0.90,
    "needs_more_pages": false
}}
```

## 识别特征
- 文档标题包含："缴税凭证"、"税款缴纳凭证"、"完税凭证"、"交税"
- 包含字段：纳税人识别号、缴款金额、缴款日期、税款所属期
- 通常由税务机关或银行出具

## 归档说明
- 纳税人作为归档目标

## 重要提示
**只输出JSON对象，不要输出任何其他文字、说明或markdown标记。直接输出纯JSON。**
"""

        return await self._parse_llm_json(prompt)

    async def _extract_audit_report(self, text: str, filename: str) -> Dict:
        """提取审计报告信息"""

        prompt = f"""你是财务审计报告信息提取专家。请提取关键信息。

## 文件名
{filename}

## 内容（前3000字符）
{text[:3000]}

## 提取任务
请提取以下关键信息，返回JSON格式：

{{
    "material_name": "审计报告标准名称（如：2024年度财务审计报告）",
    "audit_info": {{
        "report_number": "报告编号/报告文号",
        "audited_company": "被审计单位名称（完整）",
        "audited_company_credit_code": "被审计单位统一社会信用代码",
        "audit_firm": "审计机构/会计师事务所名称",
        "audit_opinion": "审计意见（无保留意见/保留意见/否定意见/无法表示意见）",
        "auditor_name": "签字注册会计师姓名",
        "report_type": "报告类型（年度审计/专项审计/离任审计等）",
        "fiscal_year": "审计年度（YYYY）"
    }},
    "company_info": {{
        "name": "被审计单位名称（归档目标）",
        "credit_code": "统一社会信用代码"
    }},
    "key_dates": {{
        "report_date": "报告日期（YYYY-MM-DD）",
        "fiscal_period_start": "会计期间起始日（YYYY-MM-DD）",
        "fiscal_period_end": "会计期间结束日（YYYY-MM-DD）"
    }},
    "confidence": 0.95,
    "needs_more_pages": false
}}

## 识别规则
1. **被审计单位**：通常在报告标题或第一段明确标注
2. **审计机构**：报告末尾盖章处，通常是"XX会计师事务所（特殊普通合伙）"
3. **审计意见**：关键词包括"我们认为"、"审计意见"、"无保留意见"等
4. **报告日期**：审计机构盖章日期
5. **会计期间**：如"2024年1月1日至2024年12月31日"

## 归档说明
- 审计报告归档到**被审计单位**名下

## 重要提示
**只输出JSON对象，不要输出任何其他文字、说明或markdown标记。直接输出纯JSON。**
"""

        return await self._parse_llm_json(prompt)

    async def _extract_generic(self, text: str, filename: str) -> Dict:
        """通用提取器（未知类型）"""

        prompt = f"""请从材料中提取关键信息。

## 文件名
{filename}

## 内容
{text[:3000]}

## 提取任务
```json
{{
    "material_name": "材料名称",
    "company_info": {{
        "name": "相关公司名称（如果有）"
    }},
    "person_info": {{
        "name": "相关人员姓名（如果有）"
    }},
    "key_info": {{
        "summary": "主要内容概述"
    }},
    "confidence": 0.50,
    "needs_more_pages": true
}}
```
"""

        return await self._parse_llm_json(prompt)

    async def _parse_llm_json(self, prompt: str) -> Dict:
        """调用LLM并解析JSON响应"""

        try:
            messages = [{"role": "user", "content": prompt}]
            response = self.llm.chat(messages)

            # 清理响应 - 移除markdown代码块标记
            cleaned_response = response.strip()
            if cleaned_response.startswith("```json"):
                cleaned_response = cleaned_response[7:]
            elif cleaned_response.startswith("```"):
                cleaned_response = cleaned_response[3:]
            if cleaned_response.endswith("```"):
                cleaned_response = cleaned_response[:-3]
            cleaned_response = cleaned_response.strip()

            # 额外清理：如果JSON后面还有其他文字（如说明、注释），只保留JSON部分
            # 找到第一个 { 和最后一个匹配的 }
            first_brace = cleaned_response.find('{')
            if first_brace != -1:
                # 从第一个{开始，找到匹配的最后一个}
                brace_count = 0
                last_brace = -1
                for i in range(first_brace, len(cleaned_response)):
                    if cleaned_response[i] == '{':
                        brace_count += 1
                    elif cleaned_response[i] == '}':
                        brace_count -= 1
                        if brace_count == 0:
                            last_brace = i
                            break

                if last_brace != -1:
                    cleaned_response = cleaned_response[first_brace:last_brace+1]

            # 解析JSON
            analysis = json.loads(cleaned_response)

            # 🔍 DEBUG: 打印关键字段
            logger.info(f"🔍 LLM提取结果预览:")
            logger.info(f"   - material_name: {analysis.get('material_name', 'N/A')[:50]}")
            logger.info(f"   - confidence: {analysis.get('confidence', 'N/A')}")
            logger.info(f"   - needs_more_pages: {analysis.get('needs_more_pages', 'N/A')}")

            # 如果是合同类型，打印甲乙方信息
            if analysis.get('contract_info'):
                party_a = analysis['contract_info'].get('party_a', {}).get('name', 'N/A')
                party_b = analysis['contract_info'].get('party_b', {}).get('name', 'N/A')
                contract_num = analysis['contract_info'].get('contract_number', 'N/A')
                logger.info(f"   - 甲方: {party_a[:30] if party_a != 'N/A' else 'N/A'}")
                logger.info(f"   - 乙方: {party_b[:30] if party_b != 'N/A' else 'N/A'}")
                logger.info(f"   - 合同编号: {contract_num}")

            return analysis

        except json.JSONDecodeError as e:
            logger.error(f"LLM返回JSON解析失败: {e}\n响应: {response[:500] if 'response' in locals() else 'N/A'}")
            return {
                "material_name": "unknown",
                "confidence": 0.0,
                "error": "JSON解析失败"
            }
        except Exception as e:
            logger.error(f"智能分析失败: {e}", exc_info=True)
            return {
                "material_name": "unknown",
                "confidence": 0.0,
                "error": str(e)
            }

    async def _old_analyze(self, extracted_content: Dict, filename: str) -> Dict:
        """智能分析材料内容"""

        text = extracted_content["text"]

        if not text or len(text.strip()) < 20:
            logger.warning(f"文本内容过少，无法进行智能分析: {filename}")
            return {
                "material_type": "unknown",
                "confidence": 0.0,
                "error": "文本内容不足"
            }

        # 构建分析prompt
        prompt = f"""你是一个专业的企业材料智能识别系统。请仔细分析材料内容，精准提取关键信息。

## 文件信息
文件名: {filename}

## 材料内容
{text[:3000]}

## 识别任务
根据材料类型，提取对应的关键字段，返回JSON格式：

```json
{{
    "material_type": "材料类型代码",
    "material_name": "材料标准名称",

    // 🔵 合同类材料专用字段（contract）
    "contract_info": {{
        "contract_number": "合同编号/协议编号",
        "party_a": {{
            "name": "甲方公司全称（完整无缩写）",
            "legal_person": "甲方法定代表人",
            "credit_code": "甲方统一社会信用代码"
        }},
        "party_b": {{
            "name": "乙方公司全称（完整无缩写）",
            "legal_person": "乙方法定代表人",
            "credit_code": "乙方统一社会信用代码"
        }},
        "amount": "合同金额（带单位，如：999900元）",
        "signing_date": "签订日期（YYYY-MM-DD）"
    }},

    // 🟢 证书/执照类材料专用字段（license/qualification/iso_cert等）
    "company_info": {{
        "name": "公司全称（完整无缩写）",
        "legal_person": "法定代表人",
        "credit_code": "统一社会信用代码（18位）",
        "address": "注册地址"
    }},

    // 🟡 证书基本信息
    "certificate_info": {{
        "certificate_number": "证书编号",
        "issuing_authority": "发证机关/颁发机构",
        "certificate_grade": "证书等级/资质等级"
    }},

    // 🟠 人员信息（身份证/学历证等）
    "person_info": {{
        "name": "姓名",
        "id_number": "身份证号",
        "education": "学历",
        "position": "职位"
    }},

    // 📅 关键日期
    "key_dates": {{
        "issue_date": "颁发日期（YYYY-MM-DD）",
        "expiry_date": "有效期至（YYYY-MM-DD）",
        "validity_period": "有效期年限"
    }},

    // 📊 质量评估
    "confidence": 0.95,
    "needs_more_pages": false,
    "notes": "补充说明"
}}
```

## 材料类型代码（material_type）
- `contract`: 合同/协议（采购合同、服务合同、框架协议等）
- `license`: 营业执照
- `qualification`: 资质证书（建筑资质、安全生产许可证等）
- `iso_cert`: ISO认证证书（ISO9001、ISO27001等）
- `education_cert`: 学历证书
- `id_card`: 身份证
- `legal_person_cert`: 法定代表人证明
- `other`: 其他类型

## 核心识别规则

### 📋 合同类型识别要点
1. **必须提取**：甲方、乙方的完整公司名称（不要缩写，如"南网数研院"应写全称）
2. **重点关注**：合同编号通常在标题下方或首页显著位置
3. **甲乙方位置**：通常在合同首页明确标注"甲方："、"乙方："
4. **金额识别**：寻找"合同金额"、"总价"、"价款"等关键词
5. **归档依据**：甲方公司作为归档目标（company_info.name填写甲方公司）

### 🏢 证书类型识别要点
1. **营业执照**：必须提取统一社会信用代码（18位）、法定代表人
2. **资质证书**：必须提取证书编号、发证机关、有效期
3. **ISO证书**：必须提取认证机构、证书编号、有效期

### 👤 人员类型识别要点
1. **身份证**：必须提取姓名和18位身份证号
2. **学历证书**：必须提取姓名、学历层次、毕业院校

## 智能分页判断（needs_more_pages）

### ✅ 可以停止扫描（false）的情况：
- **合同**：已识别到甲乙方公司名称 + 合同编号
- **营业执照**：已识别到公司名称 + 统一社会信用代码
- **身份证**：已识别到姓名 + 身份证号
- **证书（充分信息）**：证书编号、发证机关、有效期三者至少有两项

### ⚠️ 需要继续扫描（true）的情况：
- 无法确定材料类型
- 合同只看到标题，未找到甲乙方
- 证书类核心三要素（编号/机关/期限）缺失超过一半
- 公司名称只有简称，无法确定完整名称

### 核心原则
**以"能否准确归档"为标准**，而非"字段完整度"。能识别出材料类型和关联公司/人员即可。

## 输出要求
1. 严格按JSON格式输出，不要有额外的文字说明
2. 未识别的字段填null，不要臆造信息
3. 公司名称必须完整（不要使用简称或缩写）
4. 日期统一格式：YYYY-MM-DD
5. 置信度评分：0.0-1.0（核心信息齐全≥0.85，部分缺失0.60-0.84，严重缺失<0.60）
"""

        try:
            # 调用LLM
            messages = [{"role": "user", "content": prompt}]
            response = self.llm.chat(messages)

            # 清理响应 - 移除markdown代码块标记
            cleaned_response = response.strip()
            if cleaned_response.startswith("```json"):
                cleaned_response = cleaned_response[7:]  # 移除 ```json
            elif cleaned_response.startswith("```"):
                cleaned_response = cleaned_response[3:]  # 移除 ```
            if cleaned_response.endswith("```"):
                cleaned_response = cleaned_response[:-3]  # 移除结尾的 ```
            cleaned_response = cleaned_response.strip()

            # 解析JSON
            analysis = json.loads(cleaned_response)

            # 后处理
            analysis = self.post_process_analysis(analysis)
            analysis["_extracted_text"] = text  # 保存原始文本

            return analysis

        except json.JSONDecodeError as e:
            logger.error(f"LLM返回JSON解析失败: {e}\n原始响应: {response[:300]}\n清理后: {cleaned_response[:300] if 'cleaned_response' in locals() else 'N/A'}")
            return {
                "material_type": "unknown",
                "confidence": 0.0,
                "error": "分析失败"
            }
        except Exception as e:
            logger.error(f"智能分析失败: {e}", exc_info=True)
            return {
                "material_type": "unknown",
                "confidence": 0.0,
                "error": str(e)
            }

    def post_process_analysis(self, analysis: Dict) -> Dict:
        """后处理：清理和标准化"""

        # 清理公司名称
        if analysis.get("company_info", {}).get("name"):
            company_name = analysis["company_info"]["name"]
            company_name = " ".join(company_name.split())  # 去除多余空格
            company_name = company_name.replace('（', '(').replace('）', ')')  # 全角转半角
            analysis["company_info"]["name"] = company_name

        # 标准化日期格式
        if analysis.get("key_dates"):
            for date_key in ["issue_date", "expiry_date"]:
                if analysis["key_dates"].get(date_key):
                    try:
                        analysis["key_dates"][date_key] = self.parse_date(
                            analysis["key_dates"][date_key]
                        )
                    except:
                        analysis["key_dates"][date_key] = None

        # 验证统一社会信用代码格式
        if analysis.get("company_info", {}).get("credit_code"):
            credit_code = analysis["company_info"]["credit_code"]
            if not self.validate_credit_code(credit_code):
                logger.warning(f"信用代码格式可能有误: {credit_code}")
                analysis["confidence"] = max(0, analysis.get("confidence", 0) - 0.1)

        return analysis

    def parse_date(self, date_str: str) -> str:
        """解析各种日期格式为标准格式"""
        from dateutil import parser
        try:
            dt = parser.parse(date_str)
            return dt.strftime("%Y-%m-%d")
        except:
            return date_str

    def validate_credit_code(self, code: str) -> bool:
        """验证统一社会信用代码格式"""
        import re
        pattern = r'^[0-9A-HJ-NPQRTUWXY]{2}\d{6}[0-9A-HJ-NPQRTUWXY]{10}$'
        return bool(re.match(pattern, code))


class EntityMatcher:
    """智能实体匹配器"""

    async def match(self, analysis: Dict) -> Dict:
        """根据材料类型智能匹配实体（不同类型有不同归档规则）"""
        result = {
            "company_id": None,
            "company_name": None,
            "company_match_type": None,
            "person_id": None,
            "person_name": None,
            "person_match_type": None,
            "confidence": 1.0,
            "alternatives": [],
            "_material_type": None  # 保存材料类型，供后续调整置信度使用
        }

        material_type = analysis.get("material_type", "other")
        result["_material_type"] = material_type
        logger.info(f"🎯 开始实体匹配，材料类型: {material_type}")

        # 根据材料类型确定归档目标
        if material_type == "contract":
            # 合同类 - 乙方（party_b）作为归档目标（投标系统，乙方是"我们公司"）
            party_b = analysis.get("contract_info", {}).get("party_b", {})
            if party_b.get("name"):
                logger.info(f"📋 合同类：使用乙方作为归档目标 - {party_b['name']} (我们公司)")
                company_match = await self.match_company(party_b)
                result.update(company_match)

                # 🔵 合同类特殊处理：如果是新公司但合同核心信息齐全，提高实体匹配置信度
                if company_match.get("company_match_type") == "new":
                    contract_info = analysis.get("contract_info", {})
                    party_a_name = contract_info.get("party_a", {}).get("name")
                    party_b_name = contract_info.get("party_b", {}).get("name")
                    contract_number = contract_info.get("contract_number")

                    if party_a_name and party_b_name and contract_number:
                        # 三项核心信息齐全 → 提高置信度
                        result["confidence"] = 0.90  # 从0.7提高到0.90
                        logger.info(f"  ✅ 合同核心信息齐全，新公司实体匹配置信度提升至 0.90")

                # 同步到company_info
                analysis["company_info"] = party_b

        elif material_type == "project_performance":
            # 项目业绩 - 承建单位（contractor）作为归档目标
            project_info = analysis.get("project_info", {})
            contractor_name = project_info.get("contractor")
            if contractor_name:
                logger.info(f"🏗️ 项目业绩：使用承建单位作为归档目标 - {contractor_name}")
                company_match = await self.match_company({"name": contractor_name})
                result.update(company_match)
                # 同步到company_info
                if not analysis.get("company_info"):
                    analysis["company_info"] = {}
                analysis["company_info"]["name"] = contractor_name

        elif material_type == "financial_document":
            # 财务票据 - 购买方/纳税人作为归档目标
            invoice_info = analysis.get("invoice_info", {})
            tax_payment_info = analysis.get("tax_payment_info", {})

            # 优先处理发票信息
            if invoice_info and invoice_info.get("buyer"):
                buyer_name = invoice_info.get("buyer")
                buyer_tax_number = invoice_info.get("buyer_tax_number")
                logger.info(f"💰 财务票据（发票）：使用购买方作为归档目标 - {buyer_name}")
                company_match = await self.match_company({
                    "name": buyer_name,
                    "credit_code": buyer_tax_number
                })
                result.update(company_match)
                # 同步到company_info
                if not analysis.get("company_info"):
                    analysis["company_info"] = {}
                analysis["company_info"]["name"] = buyer_name
                if buyer_tax_number:
                    analysis["company_info"]["credit_code"] = buyer_tax_number
            # 处理完税证明信息
            elif tax_payment_info and tax_payment_info.get("taxpayer_name"):
                taxpayer_name = tax_payment_info.get("taxpayer_name")
                taxpayer_id = tax_payment_info.get("taxpayer_id")
                logger.info(f"💰 财务票据（完税证明）：使用纳税人作为归档目标 - {taxpayer_name}")
                company_match = await self.match_company({
                    "name": taxpayer_name,
                    "credit_code": taxpayer_id
                })
                result.update(company_match)
                # 同步到company_info
                if not analysis.get("company_info"):
                    analysis["company_info"] = {}
                analysis["company_info"]["name"] = taxpayer_name
                if taxpayer_id:
                    analysis["company_info"]["credit_code"] = taxpayer_id

        elif material_type == "tax_payment_voucher":
            # 交税凭证 - 纳税人作为归档目标
            tax_voucher_info = analysis.get("tax_voucher_info", {})
            taxpayer_name = tax_voucher_info.get("taxpayer_name")
            taxpayer_id = tax_voucher_info.get("taxpayer_id")
            if taxpayer_name:
                logger.info(f"💳 交税凭证：使用纳税人作为归档目标 - {taxpayer_name}")
                company_match = await self.match_company({
                    "name": taxpayer_name,
                    "credit_code": taxpayer_id
                })
                result.update(company_match)
                # 同步到company_info
                if not analysis.get("company_info"):
                    analysis["company_info"] = {}
                analysis["company_info"]["name"] = taxpayer_name
                if taxpayer_id:
                    analysis["company_info"]["credit_code"] = taxpayer_id

        elif material_type == "audit_report":
            # 审计报告 - 被审计单位作为归档目标
            audit_info = analysis.get("audit_info", {})
            audited_company = audit_info.get("audited_company")
            audited_company_credit_code = audit_info.get("audited_company_credit_code")
            if audited_company:
                logger.info(f"📊 审计报告：使用被审计单位作为归档目标 - {audited_company}")
                company_match = await self.match_company({
                    "name": audited_company,
                    "credit_code": audited_company_credit_code
                })
                result.update(company_match)
                # 同步到company_info
                if not analysis.get("company_info"):
                    analysis["company_info"] = {}
                analysis["company_info"]["name"] = audited_company
                if audited_company_credit_code:
                    analysis["company_info"]["credit_code"] = audited_company_credit_code

        elif material_type in ["company_qualification", "company_business", "iso_certificate"]:
            # 公司资质/工商信息/ISO认证 - 使用company_info中的公司
            company_info = analysis.get("company_info", {})
            if company_info.get("name"):
                logger.info(f"🏢 {material_type}：使用持证/获证企业作为归档目标 - {company_info['name']}")
                company_match = await self.match_company(company_info)
                result.update(company_match)

        elif material_type == "employee_document":
            # 员工证件 - 匹配人员信息
            person_info = analysis.get("person_info", {})
            if person_info.get("name"):
                logger.info(f"👤 员工证件：匹配人员 - {person_info['name']}")
                person_match = await self.match_person(person_info, result.get("company_id"))
                result.update(person_match)

        else:
            # 其他类型 - 尝试通用匹配
            logger.info(f"📄 其他类型：尝试通用匹配")
            if analysis.get("company_info", {}).get("name"):
                company_match = await self.match_company(analysis["company_info"])
                result.update(company_match)

        # 如果有人员信息且未匹配，尝试匹配人员（employee_document类型已经在上面处理了）
        if material_type != "employee_document" and analysis.get("person_info", {}).get("name"):
            person_match = await self.match_person(
                analysis["person_info"],
                result.get("company_id")
            )
            result.update(person_match)

        return result

    async def match_company(self, company_info: Dict) -> Dict:
        """匹配公司"""
        company_name = company_info.get("name")
        credit_code = company_info.get("credit_code")

        if not company_name:
            return {
                "company_id": None,
                "company_match_type": None,
                "confidence": 0.0
            }

        with get_session() as session:
            # 1. 精确匹配：统一社会信用代码
            if credit_code:
                company = session.query(Company).filter(
                    Company.credit_code == credit_code
                ).first()

                if company:
                    logger.info(f"✅ 公司精确匹配（信用代码）: {company.name}")
                    return {
                        "company_id": company.id,
                        "company_name": company.name,
                        "company_match_type": "exact_credit_code",
                        "confidence": 1.0
                    }

            # 2. 精确匹配：公司名称
            company = session.query(Company).filter(
                Company.name == company_name
            ).first()

            if company:
                logger.info(f"✅ 公司精确匹配（名称）: {company.name}")
                return {
                    "company_id": company.id,
                    "company_name": company.name,
                    "company_match_type": "exact_name",
                    "confidence": 1.0
                }

            # 3. 模糊匹配
            all_companies = session.query(Company).all()
            fuzzy_matches = []

            for comp in all_companies:
                similarity = self.calculate_company_name_similarity(
                    company_name,
                    comp.name
                )
                if similarity > 0.8:
                    fuzzy_matches.append({
                        "company_id": comp.id,
                        "company_name": comp.name,
                        "similarity": similarity
                    })

            if fuzzy_matches:
                fuzzy_matches.sort(key=lambda x: x["similarity"], reverse=True)
                best_match = fuzzy_matches[0]

                if best_match["similarity"] > 0.90:
                    logger.info(f"✅ 公司高度相似匹配: {best_match['company_name']} (相似度: {best_match['similarity']:.2f})")
                    return {
                        "company_id": best_match["company_id"],
                        "company_name": best_match["company_name"],
                        "company_match_type": "fuzzy_high",
                        "confidence": best_match["similarity"],
                        "alternatives": fuzzy_matches[1:3]
                    }
                else:
                    logger.info(f"⚠️ 公司低相似度匹配，需人工确认: {company_name}")
                    return {
                        "company_id": None,
                        "company_name": company_name,
                        "company_match_type": "fuzzy_low",
                        "confidence": 0.5,
                        "alternatives": fuzzy_matches[:3],
                        "new_company_info": company_info
                    }

            # 4. 无匹配：新公司
            logger.info(f"🆕 新公司: {company_name}")
            return {
                "company_id": None,
                "company_name": company_name,
                "company_match_type": "new",
                "confidence": 0.7,
                "new_company_info": company_info
            }

    def calculate_company_name_similarity(self, name1: str, name2: str) -> float:
        """计算公司名称相似度"""
        base_similarity = SequenceMatcher(None, name1, name2).ratio()

        # 去除公司类型后缀后的相似度
        def remove_suffix(name):
            suffixes = ["有限公司", "有限责任公司", "股份有限公司", "集团", "公司"]
            for suffix in suffixes:
                if name.endswith(suffix):
                    return name[:-len(suffix)]
            return name

        core1 = remove_suffix(name1)
        core2 = remove_suffix(name2)
        core_similarity = SequenceMatcher(None, core1, core2).ratio()

        return max(base_similarity, core_similarity)

    async def match_person(self, person_info: Dict, company_id: Optional[int]) -> Dict:
        """匹配人员"""
        person_name = person_info.get("name")
        id_number = person_info.get("id_number")

        if not person_name:
            return {"person_id": None}

        with get_session() as session:
            # 1. 身份证号精确匹配
            if id_number:
                person = session.query(Person).filter(
                    Person.id_number == id_number
                ).first()

                if person:
                    return {
                        "person_id": person.id,
                        "person_name": person.name,
                        "person_match_type": "exact_id",
                        "confidence": 1.0
                    }

            # 2. 姓名+公司匹配
            if company_id:
                person = session.query(Person).filter(
                    Person.name == person_name,
                    Person.company_id == company_id
                ).first()

                if person:
                    return {
                        "person_id": person.id,
                        "person_name": person.name,
                        "person_match_type": "name_company",
                        "confidence": 0.9
                    }

            # 3. 新人员
            return {
                "person_id": None,
                "person_name": person_name,
                "person_match_type": "new",
                "confidence": 0.6,
                "new_person_info": person_info
            }


class SmartImportPipeline:
    """智能导入流水线"""

    def __init__(self):
        self.file_processor = FileProcessor()
        self.content_extractor = ContentExtractor()
        self.intelligent_analyzer = IntelligentAnalyzer()
        self.entity_matcher = EntityMatcher()
        self.auto_archiver = AutoArchiver()

    def _update_progress(self, pending_id: int, progress_data: Dict):
        """更新处理进度到数据库"""
        if not pending_id:
            return

        import json
        try:
            with get_session() as session:
                item = session.query(PendingReview).get(pending_id)
                if item and item.status == "processing":
                    item.processing_progress = json.dumps(progress_data, ensure_ascii=False)
                    session.commit()
        except Exception as e:
            logger.warning(f"更新进度失败: {e}")

    async def process_single_file(
        self,
        file: UploadFile,
        pending_id: Optional[int] = None,
        page_numbers: Optional[List[int]] = None,
        extract_all_pages: bool = False
    ) -> Dict:
        """处理单个文件的完整流程

        参数:
        - page_numbers: 可选，指定要扫描的页码列表（从1开始），例如：[1, 3, 5]
        - extract_all_pages: 可选，是否提取所有页面为PNG（默认false）
        """

        logger.info(f"=" * 60)
        logger.info(f"📄 开始处理文件: {file.filename}")
        if page_numbers:
            logger.info(f"📄 用户指定页码: {page_numbers}")

        try:
            # 1. 保存临时文件
            self._update_progress(pending_id, {
                "stage": "saving",
                "message": "保存临时文件...",
                "current_page": 0,
                "total_pages": 0
            })

            temp_path = TEMP_DIR / f"{uuid.uuid4()}{Path(file.filename).suffix}"
            with open(temp_path, "wb") as f:
                content = await file.read()
                f.write(content)

            logger.info(f"💾 临时文件保存: {temp_path}")

            # 1.5 计算文件hash并检查重复
            self._update_progress(pending_id, {
                "stage": "checking_duplicate",
                "message": "检查文件是否已存在...",
                "current_page": 0,
                "total_pages": 0
            })

            file_hash = get_file_hash(str(temp_path))
            logger.info(f"🔍 文件Hash: {file_hash}")

            # 检查是否已导入
            with get_session() as session:
                # 先检查Material表
                existing_material = session.query(Material).filter(
                    Material.file_hash == file_hash
                ).first()

                if existing_material:
                    logger.warning(f"⚠️ 文件已存在: Material ID={existing_material.id}, {existing_material.title}")

                    # 更新pending状态为rejected
                    if pending_id:
                        try:
                            pending_item = session.query(PendingReview).get(pending_id)
                            if pending_item:
                                pending_item.status = "rejected"
                                pending_item.review_notes = f"文件已存在(Material ID={existing_material.id})"
                                pending_item.processing_progress = json.dumps({
                                    "stage": "completed",
                                    "message": "文件已存在，已自动拒绝"
                                }, ensure_ascii=False)
                                session.commit()
                        except Exception as e:
                            logger.warning(f"更新pending状态失败: {e}")

                    # 清理临时文件
                    try:
                        Path(temp_path).unlink()
                    except:
                        pass

                    return {
                        "status": "duplicate",
                        "filename": file.filename,
                        "message": "该文件已导入，请勿重复上传",
                        "existing_material": {
                            "id": existing_material.id,
                            "title": existing_material.title,
                            "section": existing_material.section,
                            "company_id": existing_material.company_id,
                            "created_at": existing_material.created_at.isoformat() if existing_material.created_at else None
                        }
                    }

                # 再检查PendingReview表（待审核）
                existing_pending = session.query(PendingReview).filter(
                    PendingReview.file_hash == file_hash,
                    PendingReview.status.in_(['pending', 'processing'])
                ).first()

                if existing_pending:
                    logger.warning(f"⚠️ 文件正在待审核: Pending ID={existing_pending.id}")

                    # 更新当前pending状态为rejected（如果不是同一个）
                    if pending_id and pending_id != existing_pending.id:
                        try:
                            pending_item = session.query(PendingReview).get(pending_id)
                            if pending_item:
                                pending_item.status = "rejected"
                                pending_item.review_notes = f"文件已在待审核列表中(Pending ID={existing_pending.id})"
                                pending_item.processing_progress = json.dumps({
                                    "stage": "completed",
                                    "message": "文件已在待审核中，已自动拒绝"
                                }, ensure_ascii=False)
                                session.commit()
                        except Exception as e:
                            logger.warning(f"更新pending状态失败: {e}")

                    # 清理临时文件
                    try:
                        Path(temp_path).unlink()
                    except:
                        pass

                    return {
                        "status": "duplicate",
                        "filename": file.filename,
                        "message": "该文件正在待审核中，请勿重复上传",
                        "existing_pending": {
                            "id": existing_pending.id,
                            "filename": existing_pending.filename,
                            "status": existing_pending.status,
                            "created_at": existing_pending.created_at.isoformat() if existing_pending.created_at else None
                        }
                    }

            # 2. 格式识别
            self._update_progress(pending_id, {
                "stage": "analyzing",
                "message": "分析文件格式...",
                "current_page": 0,
                "total_pages": 0
            })

            file_info = self.file_processor.analyze_file(file)
            logger.info(f"📋 文件类型: {file_info['type']}, 提示: {file_info['hints']}")

            # 3. 内容提取
            if page_numbers:
                # 用户指定页码 - 直接扫描指定页面
                self._update_progress(pending_id, {
                    "stage": "extracting",
                    "message": f"扫描用户指定的页码: {page_numbers}...",
                    "current_page": 0,
                    "total_pages": 0
                })

                extracted_content = await self.content_extractor.extract(
                    str(temp_path),
                    file_info,
                    progress_callback=lambda progress: self._update_progress(pending_id, progress),
                    page_numbers=page_numbers
                )
                logger.info(f"📝 用户指定页码提取完成: {len(extracted_content['text'])} 字符, {len(extracted_content.get('images', []))} 张图片")
            else:
                # 默认模式 - 第一阶段：先扫描前5页
                self._update_progress(pending_id, {
                    "stage": "extracting",
                    "message": "第一阶段：提取前5页内容...",
                    "current_page": 0,
                    "total_pages": 0
                })

                # 先扫描前5页
                extracted_content = await self.content_extractor.extract(
                    str(temp_path),
                    file_info,
                    progress_callback=lambda progress: self._update_progress(pending_id, progress),
                    start_page=0,
                    page_limit=5
                )
                logger.info(f"📝 第一阶段提取完成: {len(extracted_content['text'])} 字符, {len(extracted_content.get('images', []))} 张图片")

            # 4. 智能分析
            if page_numbers:
                # 用户指定页码 - 直接分析，不分阶段
                progress_data = {
                    "stage": "analyzing",
                    "message": f"LLM分析中（基于指定页码: {page_numbers}）...",
                    "current_page": 0,
                    "total_pages": extracted_content.get('metadata', {}).get('page_count', 0)
                }
                # 如果有OCR结果，包含进去
                if 'ocr_results' in extracted_content:
                    progress_data['ocr_results'] = extracted_content['ocr_results']

                self._update_progress(pending_id, progress_data)

                analysis = await self.intelligent_analyzer.analyze(extracted_content, file.filename)
                logger.info(f"🤖 用户指定页码分析: 类型={analysis.get('material_type')}, 置信度={analysis.get('confidence', 0):.2f}")

            else:
                # 默认模式 - 分阶段分析
                # 更新进度 - 显示第一阶段OCR结果
                progress_data = {
                    "stage": "analyzing_phase1",
                    "message": "LLM分析中（基于前5页）...",
                    "current_page": 5,
                    "total_pages": extracted_content.get('metadata', {}).get('page_count', 0)
                }
                # 如果有OCR结果，包含进去
                if 'ocr_results' in extracted_content:
                    progress_data['ocr_results'] = extracted_content['ocr_results']

                self._update_progress(pending_id, progress_data)

                analysis = await self.intelligent_analyzer.analyze(extracted_content, file.filename)
                logger.info(f"🤖 第一阶段分析: 类型={analysis.get('material_type')}, 置信度={analysis.get('confidence', 0):.2f}")

                # 5. 判断是否需要扫描更多页面
                needs_more_pages = analysis.get('needs_more_pages', False)
                total_doc_pages = extracted_content.get('metadata', {}).get('page_count', 0)

            if not page_numbers and needs_more_pages and total_doc_pages > 5:
                # 第二阶段：继续扫描剩余页面
                logger.info(f"⚠️ LLM判断信息不足，继续扫描剩余页面（第6-10页）...")

                self._update_progress(pending_id, {
                    "stage": "extracting_phase2",
                    "message": "第二阶段：继续扫描第6-10页...",
                    "current_page": 5,
                    "total_pages": total_doc_pages
                })

                # 扫描第6-10页（索引5-9）
                additional_content = await self.content_extractor.extract(
                    str(temp_path),
                    file_info,
                    progress_callback=lambda progress: self._update_progress(pending_id, progress),
                    start_page=5,
                    page_limit=5
                )

                # 合并内容
                extracted_content['text'] += "\n\n" + additional_content['text']
                extracted_content['images'].extend(additional_content.get('images', []))

                # 合并OCR结果（如果有的话）
                if 'ocr_results' in additional_content:
                    if 'ocr_results' not in extracted_content:
                        extracted_content['ocr_results'] = []
                    extracted_content['ocr_results'].extend(additional_content['ocr_results'])

                logger.info(f"📝 第二阶段提取完成，累计: {len(extracted_content['text'])} 字符，共{len(extracted_content.get('ocr_results', []))}页OCR结果")

                # 更新进度 - 显示完整OCR结果
                progress_data = {
                    "stage": "analyzing_phase2",
                    "message": "LLM重新分析中（基于完整内容）...",
                    "current_page": 10,
                    "total_pages": total_doc_pages
                }
                # 包含完整的OCR结果
                if 'ocr_results' in extracted_content:
                    progress_data['ocr_results'] = extracted_content['ocr_results']

                self._update_progress(pending_id, progress_data)

                analysis = await self.intelligent_analyzer.analyze(extracted_content, file.filename)
                logger.info(f"🤖 第二阶段分析: 类型={analysis.get('material_type')}, 置信度={analysis.get('confidence', 0):.2f}")

            # 5. 实体匹配
            entities = await self.entity_matcher.match(analysis)
            logger.info(f"🔗 实体匹配: 公司={entities.get('company_name')}, 匹配类型={entities.get('company_match_type')}")

            # 5.5 对于合同类型，智能识别关键页面（首页、金额/服务范围页、签字页）
            if analysis.get('material_type') == 'contract' and str(temp_path).lower().endswith('.pdf') and not page_numbers:
                logger.info("📋 检测到合同类型，开始识别关键页面...")

                import fitz
                pdf_doc = fitz.open(str(temp_path))
                total_pdf_pages = len(pdf_doc)

                # 先尝试按页提取原生文字
                native_text_list = []
                for page_idx in range(total_pdf_pages):
                    page_text = pdf_doc[page_idx].get_text().strip()
                    if page_text:
                        native_text_list.append({
                            "page": page_idx + 1,
                            "text": page_text
                        })
                pdf_doc.close()

                total_native_chars = sum(len(item['text']) for item in native_text_list)

                if total_native_chars >= 500:
                    # 非扫描件：有足够的原生文字，直接用于关键页面分析
                    logger.info(f"📄 非扫描件PDF，共{total_pdf_pages}页，{len(native_text_list)}页有文字，共{total_native_chars}字符")
                    ocr_text_list = native_text_list
                else:
                    # 扫描件：依赖OCR文本
                    ocr_results = extracted_content.get('ocr_results', [])
                    scanned_pages = len(ocr_results)

                    if scanned_pages < total_pdf_pages:
                        logger.info(f"📄 合同共{total_pdf_pages}页，已扫描{scanned_pages}页，继续扫描剩余页面...")

                        self._update_progress(pending_id, {
                            "stage": "scanning_all_pages",
                            "message": f"扫描合同全部页面（{scanned_pages+1}-{total_pdf_pages}页）...",
                            "current_page": scanned_pages,
                            "total_pages": total_pdf_pages
                        })

                        remaining_content = await self.content_extractor.extract(
                            str(temp_path),
                            file_info,
                            progress_callback=lambda progress: self._update_progress(pending_id, progress),
                            start_page=scanned_pages,
                            page_limit=total_pdf_pages - scanned_pages
                        )

                        if 'ocr_results' not in extracted_content:
                            extracted_content['ocr_results'] = []
                        if 'ocr_results' in remaining_content:
                            extracted_content['ocr_results'].extend(remaining_content['ocr_results'])

                        logger.info(f"✅ 全部页面扫描完成，共{len(extracted_content.get('ocr_results', []))}页")

                    # 从OCR缓存获取文本
                    ocr_text_list = []
                    logger.info(f"  📄 从缓存读取{len(extracted_content.get('ocr_results', []))}页的完整OCR文本...")

                    for ocr_item in extracted_content.get('ocr_results', []):
                        page_display = ocr_item.get("page", 0)
                        page_num = page_display - 1

                        cached_ocr = get_cached_ocr(str(temp_path), page_num)
                        if cached_ocr:
                            full_text = cached_ocr.get("text", "")
                        else:
                            full_text = ocr_item.get("preview", "")
                            logger.warning(f"    ⚠️ 第{page_display}页缓存未命中，使用preview: {len(full_text)}字符")

                        ocr_text_list.append({
                            "page": page_display,
                            "text": full_text
                        })

                    logger.info(f"  ✅ 已准备{len(ocr_text_list)}页文本，总字符数: {sum(len(item['text']) for item in ocr_text_list)}")

                # 使用LLM分析识别关键页面
                if ocr_text_list:
                    self._update_progress(pending_id, {
                        "stage": "identifying_key_pages",
                        "message": "智能识别关键页面（首页、金额、签字页）...",
                        "current_page": total_pdf_pages,
                        "total_pages": total_pdf_pages
                    })

                    key_pages_result = await analyze_contract_key_pages(ocr_text_list, total_pdf_pages)
                    analysis['contract_key_pages'] = key_pages_result.get('key_pages', [])
                    analysis['key_pages_summary'] = key_pages_result.get('analysis_summary', '')
                    logger.info(f"✅ 关键页面识别完成: {key_pages_result.get('analysis_summary', '')}")
                else:
                    logger.warning("⚠️ 无文本内容，跳过关键页面识别，将使用默认策略（首页+末页）")

            # 6. 计算总体置信度
            overall_confidence = self.calculate_confidence(analysis, entities)
            logger.info(f"📊 总体置信度: {overall_confidence:.2f}")

            # 6.5 检测证书版本更新（仅针对证书类材料）
            version_info = None
            material_type = analysis.get('material_type')
            if material_type in ['iso_certificate', 'company_qualification', 'company_business']:
                company_id = entities.get('company_id')
                if company_id:
                    with get_session() as session:
                        # 查找同类证书
                        existing_cert = find_existing_certificate(
                            company_id=company_id,
                            material_type=material_type,
                            analysis=analysis,
                            session=session
                        )

                        if existing_cert:
                            # 判断是否为更新版本
                            is_newer, confidence, reason = is_newer_version(analysis, existing_cert)

                            if is_newer:
                                logger.info(f"🔄 检测到证书更新: {reason}, 置信度={confidence:.2f}")
                                version_info = {
                                    "is_update": True,
                                    "old_material_id": existing_cert.id,
                                    "old_material_title": existing_cert.title,
                                    "update_confidence": confidence,
                                    "update_reason": reason
                                }
                                # 保存到analysis中，供归档时使用
                                analysis['_version_info'] = version_info
                            else:
                                # 判断为旧版本或重复，降低置信度，避免自动归档
                                logger.warning(f"⚠️ {reason}")
                                if confidence >= 0.8:
                                    # 高置信度判断为旧版本/重复，直接拒绝
                                    logger.error(f"❌ 检测到重复或旧版本证书，拒绝导入")

                                    # 更新pending状态为rejected
                                    if pending_id:
                                        try:
                                            with get_session() as session:
                                                pending_item = session.query(PendingReview).get(pending_id)
                                                if pending_item:
                                                    pending_item.status = "rejected"
                                                    pending_item.review_notes = f"检测到旧版本或重复证书(Material ID={existing_cert.id})"
                                                    pending_item.processing_progress = json.dumps({
                                                        "stage": "completed",
                                                        "message": f"检测到旧版本证书：{reason}"
                                                    }, ensure_ascii=False)
                                                    session.commit()
                                        except Exception as e:
                                            logger.warning(f"更新pending状态失败: {e}")

                                    # 清理临时文件
                                    try:
                                        Path(temp_path).unlink()
                                    except:
                                        pass

                                    return {
                                        "status": "duplicate_old_version",
                                        "filename": file.filename,
                                        "message": f"检测到旧版本或重复证书：{reason}",
                                        "existing_certificate": {
                                            "id": existing_cert.id,
                                            "title": existing_cert.title,
                                            "expiry_date": existing_cert.expiry_date.isoformat() if existing_cert.expiry_date else None
                                        }
                                    }

            # 7. 决策：自动归档 or 人工审核
            if overall_confidence >= 0.85:
                # 自动归档
                logger.info(f"✅ 置信度高，执行自动归档")
                material = await self.auto_archiver.archive(
                    temp_path=str(temp_path),
                    filename=file.filename,
                    file_info=file_info,
                    analysis=analysis,
                    entities=entities,
                    page_numbers=page_numbers,
                    extract_all_pages=extract_all_pages
                )

                # 如果有pending_id，更新其状态
                if pending_id:
                    with get_session() as session:
                        item = session.query(PendingReview).get(pending_id)
                        if item:
                            item.status = "approved"
                            item.material_id = material["id"]
                            item.confidence = overall_confidence
                            item.processing_progress = None
                            session.commit()

                return {
                    "status": "auto_archived",
                    "material_id": material["id"],
                    "filename": file.filename,
                    "confidence": overall_confidence,
                    "message": f"已自动归档: {material['title']}"
                }
            else:
                # 如果已有pending_id，更新它；否则创建新的
                if pending_id:
                    # 更新现有记录
                    with get_session() as session:
                        item = session.query(PendingReview).get(pending_id)
                        if item:
                            item.analysis_json = json.dumps(analysis, ensure_ascii=False)
                            item.entities_json = json.dumps(entities, ensure_ascii=False)
                            item.confidence = overall_confidence
                            item.status = "pending"
                            item.processing_progress = None
                            session.commit()
                    result_pending_id = pending_id
                    logger.info(f"⚠️ 置信度不足，已更新待审核记录 (ID: {pending_id})")
                else:
                    # 创建新的待审核项
                    result_pending_id = await self.create_pending_review(
                        temp_path=str(temp_path),
                        filename=file.filename,
                        file_info=file_info,
                        analysis=analysis,
                        entities=entities,
                        confidence=overall_confidence,
                        file_hash=file_hash
                    )
                    logger.info(f"⚠️ 置信度不足，已加入待审核队列 (ID: {result_pending_id})")

                return {
                    "status": "pending_review",
                    "pending_id": result_pending_id,
                    "filename": file.filename,
                    "confidence": overall_confidence,
                    "message": "置信度不足，需人工审核"
                }

        except Exception as e:
            logger.error(f"❌ 处理失败: {file.filename} - {e}", exc_info=True)
            return {
                "status": "failed",
                "filename": file.filename,
                "error": str(e)
            }

    def calculate_confidence(self, analysis: Dict, entities: Dict) -> float:
        """计算总体置信度（根据材料类型智能调整）"""
        analysis_conf = analysis.get("confidence", 0)
        entity_conf = entities.get("confidence", 0)
        material_type = analysis.get("material_type", "other")

        # 综合评分
        overall = (analysis_conf * 0.6 + entity_conf * 0.4)

        # 调整因子（根据实体匹配类型）
        company_match_type = entities.get("company_match_type")

        if company_match_type == "exact_credit_code":
            # 统一社会信用代码精确匹配 → 高置信度
            overall += 0.1

        elif company_match_type == "new":
            # 新公司 → 根据材料类型区别对待

            if material_type == "contract":
                # 🔵 合同类：如果合同核心信息齐全，新公司不应降低置信度
                contract_info = analysis.get("contract_info", {})
                party_a_name = contract_info.get("party_a", {}).get("name")
                party_b_name = contract_info.get("party_b", {}).get("name")
                contract_number = contract_info.get("contract_number")

                if party_a_name and party_b_name and contract_number:
                    # 三项核心信息齐全 → 不惩罚新公司
                    logger.info(f"  ✅ 合同核心信息齐全（甲乙方+编号），新公司不降低置信度")
                    overall -= 0.0  # 不惩罚
                else:
                    # 信息不全 → 适度惩罚
                    overall -= 0.10

            elif material_type in ["company_business", "company_qualification", "iso_certificate"]:
                # 🟢 公司证书类：新公司可能是识别错误 → 适度惩罚
                overall -= 0.10

            elif material_type == "employee_document":
                # 🟡 员工证件：新公司影响不大（主要看人员信息）
                overall -= 0.05

            else:
                # 其他类型：标准惩罚
                overall -= 0.15

        return min(1.0, max(0.0, overall))

    async def create_pending_review(
        self,
        temp_path: str,
        filename: str,
        file_info: Dict,
        analysis: Dict,
        entities: Dict,
        confidence: float,
        file_hash: str = None
    ) -> int:
        """创建待审核项"""
        with get_session() as session:
            pending = PendingReview(
                file_path=temp_path,
                filename=filename,
                file_type=file_info["type"],
                file_hash=file_hash,
                analysis_json=json.dumps(analysis, ensure_ascii=False),
                entities_json=json.dumps(entities, ensure_ascii=False),
                version_info_json=json.dumps({"is_update": False}, ensure_ascii=False),
                confidence=int(confidence * 100),
                status="pending"
            )
            session.add(pending)
            session.commit()
            return pending.id


class AutoArchiver:
    """自动归档器 - 创建材料记录和实体"""

    def __init__(self):
        self.data_dir = Path(os.getenv("DATA_DIR", "data"))
        self.files_dir = self.data_dir / "files"
        self.images_dir = self.data_dir / "images"
        self.files_dir.mkdir(parents=True, exist_ok=True)
        self.images_dir.mkdir(parents=True, exist_ok=True)

    async def archive(
        self,
        temp_path: str,
        filename: str,
        file_info: Dict,
        analysis: Dict,
        entities: Dict,
        page_numbers: Optional[List[int]] = None,
        extract_all_pages: bool = False
    ) -> Dict:
        """执行自动归档

        Args:
            page_numbers: 用户手动选择的页码（从1开始），这些页面会全部提取为PNG
            extract_all_pages: 是否提取所有页面为PNG（手动选页模式下批准时的选项）
        """
        logger.info(f"🗄️ 开始自动归档: {filename}")
        if extract_all_pages:
            logger.info(f"  📄 用户选择：提取所有页面为PNG")

        # 计算文件hash
        file_hash = get_file_hash(temp_path)

        with get_session() as session:
            # 1. 创建或获取公司
            company_id = await self._get_or_create_company(session, entities, analysis)

            # 2. 创建或获取人员
            person_id = await self._get_or_create_person(session, entities, analysis, company_id)

            # 3. 保存文件到永久位置
            saved_file_path = self._save_file_permanent(temp_path, filename, file_info)

            # 4. 创建Document记录（如果需要）
            document = self._get_or_create_document(session, filename, company_id)

            # 5. 创建Material记录（如果提取所有页面为PNG，则跳过创建PDF主记录）
            material = None
            if not extract_all_pages:
                # 正常模式：创建主Material记录指向PDF
                material = self._create_material(
                    session,
                    document_id=document.id,
                    company_id=company_id,
                    person_id=person_id,
                    saved_file_path=saved_file_path,
                    filename=filename,
                    analysis=analysis,
                    file_hash=file_hash
                )

            # 6. 智能提取关键页面PNG（根据材料类型）
            first_page_png = None  # 记录第一张PNG，用于返回主material ID

            if saved_file_path.lower().endswith(".pdf"):
                material_type = analysis.get("material_type", "other")
                material_name = analysis.get("material_name")

                try:
                    import fitz
                    pdf_doc = fitz.open(saved_file_path)
                    total_pages = len(pdf_doc)
                    pdf_doc.close()

                    pages_to_extract = []

                    if extract_all_pages:
                        # 提取所有页面为PNG（用户勾选了"提取所有页面"选项）
                        section_name = material_name if material_name else filename
                        logger.info(f"  📄 提取所有页面模式: 共 {total_pages} 页")
                        for pn in range(1, total_pages + 1):
                            pages_to_extract.append({
                                "page_num": pn - 1,  # 转为0-based
                                "section": section_name,
                                "material_type": f"page_{pn}",
                                "title_suffix": f"第{pn}页"
                            })
                    elif page_numbers:
                        # 手动选页模式：只提取用户选的页面为PNG
                        section_name = material_name if material_name else filename
                        for pn in sorted(page_numbers):
                            if 1 <= pn <= total_pages:
                                pages_to_extract.append({
                                    "page_num": pn - 1,  # 转为0-based
                                    "section": section_name,
                                    "material_type": f"contract_page_{pn}",
                                    "title_suffix": f"第{pn}页"
                                })
                        logger.info(f"  📄 手动选页模式: 提取用户选择的 {len(pages_to_extract)} 页")

                    elif material_type == "contract" and analysis.get("contract_key_pages"):
                        # LLM识别的合同关键页面（非空时才使用）
                        logger.info(f"  📄 使用LLM识别的合同关键页面...")
                        key_pages = analysis["contract_key_pages"]

                        for key_page in key_pages:
                            pages_to_extract.append({
                                "page_num": key_page.get("page_num", 0),
                                "section": material_name if material_name else "合同",
                                "material_type": f"contract_{key_page.get('page_type', 'page')}",
                                "title_suffix": key_page.get("title_suffix", "页面")
                            })

                        logger.info(f"  📋 LLM识别: {analysis.get('key_pages_summary', '')}")

                    else:
                        # 其他材料类型或LLM未识别关键页面，使用默认策略
                        pages_to_extract = get_pages_to_extract(material_type, total_pages, material_name)

                    if pages_to_extract:
                        logger.info(f"  📄 {material_type} - 开始提取关键页面 ({len(pages_to_extract)}页)...")

                        for i, page_config in enumerate(pages_to_extract):
                            page_png = extract_pdf_page_to_png(
                                saved_file_path,
                                page_num=page_config["page_num"],
                                output_dir=self.images_dir,
                                dpi=300
                            )

                            if page_png:
                                # 提取所有页面模式或手动选页模式：所有PNG都创建为独立记录
                                if extract_all_pages or page_numbers:
                                    page_material = Material(
                                        document_id=document.id,
                                        company_id=company_id,
                                        person_id=person_id,
                                        title=f"{analysis.get('material_name', filename)} - {page_config['title_suffix']}",
                                        section=page_config["section"],
                                        image_filename=Path(page_png).name,
                                        image_path=page_png,
                                        file_size=Path(page_png).stat().st_size,
                                        material_type=page_config["material_type"],
                                        ocr_status="completed",
                                        ocr_processed_at=datetime.utcnow()
                                    )
                                    session.add(page_material)
                                    session.flush()  # 获取ID

                                    # 记录第一个PNG作为主记录（用于返回material_id）
                                    if i == 0:
                                        material = page_material
                                        if extract_all_pages:
                                            logger.info(f"  📄 提取所有页面模式：只创建PNG，不保留PDF记录")

                                    logger.info(f"  ✓ 创建{page_config['section']}: {page_material.title}")

                                elif i == 0 and material:
                                    # 自动模式：第一张PNG更新主material记录
                                    material.title = f"{analysis.get('material_name', filename)} - {page_config['title_suffix']}"
                                    material.image_filename = Path(page_png).name
                                    material.image_path = page_png
                                    material.file_size = Path(page_png).stat().st_size
                                    material.material_type = page_config["material_type"]
                                    logger.info(f"  ✓ 主记录更新为: {material.title}")
                                else:
                                    # 后续PNG：创建新的material记录
                                    page_material = Material(
                                        document_id=document.id,
                                        company_id=company_id,
                                        person_id=person_id,
                                        title=f"{analysis.get('material_name', filename)} - {page_config['title_suffix']}",
                                        section=page_config["section"],
                                        image_filename=Path(page_png).name,
                                        image_path=page_png,
                                        file_size=Path(page_png).stat().st_size,
                                        material_type=page_config["material_type"],
                                        ocr_status="completed",
                                        ocr_processed_at=datetime.utcnow()
                                    )
                                    session.add(page_material)
                                    logger.info(f"  ✓ 创建{page_config['section']}: {page_material.title}")

                except Exception as e:
                    logger.warning(f"  ⚠️ 提取页面失败: {e}")

            # 7. 处理证书版本替换
            version_replaced = False
            if '_version_info' in analysis and analysis['_version_info'].get('is_update'):
                version_info = analysis['_version_info']
                update_confidence = version_info.get('update_confidence', 0)

                # 置信度 >= 0.8 自动替换
                if update_confidence >= 0.8:
                    logger.info(f"🔄 自动替换旧版本证书 (置信度={update_confidence:.2f})")

                    from certificate_matcher import replace_with_newer_version

                    success = await replace_with_newer_version(
                        old_material=session.query(Material).get(version_info['old_material_id']),
                        new_material_id=material.id,
                        reason=version_info['update_reason'],
                        session=session
                    )

                    if success:
                        version_replaced = True
                        logger.info(f"  ✅ 证书版本替换成功: {version_info['old_material_title']} → {material.title}")
                    else:
                        logger.warning(f"  ⚠️ 证书版本替换失败")
                else:
                    logger.info(f"⚠️ 检测到可能的证书更新，但置信度较低({update_confidence:.2f})，需要人工确认")

            session.commit()

            # 在session关闭前获取需要的数据
            material_id = material.id
            material_title = material.title

            logger_msg = f"✅ 归档完成: Material ID={material_id}, {material_title}"
            if version_replaced:
                logger_msg += " (已替换旧版本)"
            logger.info(logger_msg)

        # 6. 清理临时文件（移到session外）
        try:
            Path(temp_path).unlink()
        except:
            pass

        # 返回包含material信息的字典，而不是ORM对象
        return {
            "id": material_id,
            "title": material_title
        }

    async def _get_or_create_company(
        self,
        session,
        entities: Dict,
        analysis: Dict
    ) -> Optional[int]:
        """获取或创建公司"""
        company_id = entities.get("company_id")

        # 如果已匹配到公司，直接返回
        if company_id:
            logger.info(f"  ✓ 使用已匹配公司 ID={company_id}")
            return company_id

        # 如果需要创建新公司
        if entities.get("new_company_info"):
            company_info = entities["new_company_info"]
            company = Company(
                name=company_info.get("name"),
                legal_person=company_info.get("legal_person"),
                credit_code=company_info.get("credit_code"),
                address=company_info.get("address")
            )
            session.add(company)
            session.flush()
            logger.info(f"  ✓ 创建新公司: {company.name} (ID={company.id})")
            return company.id

        return None

    async def _get_or_create_person(
        self,
        session,
        entities: Dict,
        analysis: Dict,
        company_id: Optional[int]
    ) -> Optional[int]:
        """获取或创建人员"""
        person_id = entities.get("person_id")

        # 如果已匹配到人员，直接返回
        if person_id:
            logger.info(f"  ✓ 使用已匹配人员 ID={person_id}")
            return person_id

        # 如果需要创建新人员
        if entities.get("new_person_info"):
            person_info = entities["new_person_info"]
            person = Person(
                name=person_info.get("name"),
                id_number=person_info.get("id_number"),
                education=person_info.get("education"),
                position=person_info.get("position"),
                company_id=company_id
            )
            session.add(person)
            session.flush()
            logger.info(f"  ✓ 创建新人员: {person.name} (ID={person.id})")
            return person.id

        return None

    def _save_file_permanent(
        self,
        temp_path: str,
        filename: str,
        file_info: Dict
    ) -> str:
        """保存文件到永久位置"""
        # 生成唯一文件名
        ext = Path(filename).suffix
        unique_filename = f"{uuid.uuid4()}{ext}"

        # 根据文件类型选择目录
        if file_info["type"] == "image":
            target_dir = self.images_dir
        else:
            target_dir = self.files_dir

        target_path = target_dir / unique_filename

        # 移动文件
        import shutil
        shutil.move(temp_path, target_path)

        logger.info(f"  ✓ 文件已保存: {target_path}")
        return str(target_path)

    def _get_or_create_document(
        self,
        session,
        filename: str,
        company_id: Optional[int]
    ) -> Document:
        """获取或创建Document记录"""
        # 对于智能导入，每个文件创建一个独立的Document
        document = Document(
            filename=filename,
            company_id=company_id,
            section_count=0,
            image_count=1
        )
        session.add(document)
        session.flush()
        logger.info(f"  ✓ 创建Document记录 ID={document.id}")
        return document

    def _create_material(
        self,
        session,
        document_id: int,
        company_id: Optional[int],
        person_id: Optional[int],
        saved_file_path: str,
        filename: str,
        analysis: Dict,
        file_hash: str = None,
        override_section: str = None
    ) -> Material:
        """创建Material记录"""
        # 解析有效期
        expiry_date = None
        if analysis.get("key_dates", {}).get("expiry_date"):
            try:
                expiry_date = datetime.strptime(
                    analysis["key_dates"]["expiry_date"],
                    "%Y-%m-%d"
                ).date()
            except:
                pass

        # 确定section：优先使用override_section，否则合同类型用材料名称（完整合同名），其他类型用通用分类
        if override_section:
            section = override_section
        else:
            material_type = analysis.get("material_type")
            if material_type == "contract" and analysis.get("material_name"):
                section = analysis["material_name"]
            else:
                section = self._determine_section(material_type)

        # 如果没有传入file_hash，计算它
        if not file_hash:
            file_hash = get_file_hash(saved_file_path)

        # 创建材料
        material = Material(
            document_id=document_id,
            company_id=company_id,
            person_id=person_id,
            title=analysis.get("material_name") or filename,
            section=section,
            image_filename=Path(saved_file_path).name,
            image_path=saved_file_path,
            file_size=Path(saved_file_path).stat().st_size,
            file_hash=file_hash,
            material_type=analysis.get("material_type"),
            ocr_text=analysis.get("_extracted_text"),
            extracted_json=json.dumps({
                k: v for k, v in analysis.items()
                if k != "_extracted_text"
            }, ensure_ascii=False),
            expiry_date=expiry_date,
            ocr_status="completed",
            ocr_processed_at=datetime.utcnow()
        )

        session.add(material)
        session.flush()

        logger.info(f"  ✓ 创建Material记录: {material.title} (ID={material.id})")
        return material

    def _determine_section(self, material_type: Optional[str]) -> str:
        """根据材料类型确定section"""
        type_to_section = {
            "license": "营业执照",
            "qualification": "资质证书",
            "iso_cert": "ISO认证",
            "id_card": "身份证",
            "education_cert": "学历证书",
            "legal_person_cert": "法人证明",
            "contract": "合同",
            "tax_payment_cert": "完税证明",
            "tax_payment_voucher": "交税凭证",
            "audit_report": "审计报告",
        }
        return type_to_section.get(material_type, "其他材料")
