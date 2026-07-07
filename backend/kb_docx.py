"""
Word Document (.docx/.doc) Text Extraction.

Extracts text content from Word documents using python-docx.
Handles paragraphs, tables, and headers/footers.
"""

import os
import logging
from typing import Optional

logger = logging.getLogger("materialhub.kb_docx")

MAX_TEXT_LENGTH = 50000  # Cap extracted text to avoid memory issues


def extract_text_from_docx(file_path: str) -> str:
    """Extract full text content from a .docx file.

    Extracts:
    - All paragraphs (with style info for heading detection)
    - All table cells
    - Headers and footers

    Args:
        file_path: path to .docx file

    Returns:
        Extracted text string (empty on failure)
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""

    if not os.path.exists(file_path):
        logger.warning("Word file not found: %s", file_path)
        return ""

    try:
        doc = Document(file_path)
        parts = []

        # 1. Extract paragraphs with heading detection
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower() if para.style else ""

            # Mark headings for better chunking later
            if "heading" in style_name or "title" in style_name:
                parts.append(f"\n## {text}\n")
            else:
                parts.append(text)

        # 2. Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        # 3. Extract headers and footers
        try:
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        for para in header.paragraphs:
                            if para.text.strip():
                                parts.append(f"[页眉] {para.text.strip()}")
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        for para in footer.paragraphs:
                            if para.text.strip():
                                parts.append(f"[页脚] {para.text.strip()}")
        except Exception:
            # Headers/footers extraction is best-effort
            pass

        full_text = "\n".join(parts)

        # Cap length
        if len(full_text) > MAX_TEXT_LENGTH:
            full_text = full_text[:MAX_TEXT_LENGTH]
            logger.info("Word text truncated to %d chars", MAX_TEXT_LENGTH)

        logger.info("Extracted %d chars from %s", len(full_text), os.path.basename(file_path))
        return full_text.strip()

    except Exception as e:
        logger.error("Failed to extract text from %s: %s", file_path, e)
        return ""


def extract_text_from_word(file_path: str) -> str:
    """Extract text from Word document (.docx or legacy .doc).

    For .docx: uses python-docx (native, fast)
    For .doc (legacy): tries mammoth conversion, falls back to empty

    Args:
        file_path: path to Word file

    Returns:
        Extracted text string
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        return extract_text_from_docx(file_path)

    if ext == ".doc":
        # Legacy .doc format — try mammoth
        try:
            import mammoth
            with open(file_path, "rb") as f:
                result = mammoth.extract_raw_text(f)
                text = result.value.strip()
                logger.info("Extracted %d chars from legacy .doc via mammoth", len(text))
                return text[:MAX_TEXT_LENGTH] if len(text) > MAX_TEXT_LENGTH else text
        except Exception as e:
            logger.warning("Legacy .doc extraction failed for %s: %s", file_path, e)
            return ""

    # Unknown extension — try docx anyway
    return extract_text_from_docx(file_path)
